
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Add a centered title
title = doc.add_paragraph("Dictate Mode Document", style='Title')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a placeholder paragraph
doc.add_paragraph("Start dictating your content here...", style='Normal')

# Add a section title for voice typing tips
tips_title = doc.add_paragraph("Voice Typing Tips", style='Heading 1')
tips_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add bullet points with practical tips
tips = [
    "Speak clearly and at a steady pace.",
    "Use voice commands for punctuation, such as 'comma', 'period', or 'question mark'.",
    "Review and edit the text after dictation to correct any errors.",
    "Use formatting commands like 'new line' or 'new paragraph'.",
    "Avoid background noise for better accuracy."
]

for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

# Save the document
doc.save("dictate_mode.docx")
